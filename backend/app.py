from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import logging

app = Flask(__name__)

# Configurar logging para evitar problemas con WSGI servers
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@app.before_request
def log_request_info():
    try:
        logger.info(f"SOLICITUD ENTRANTE: {request.method} {request.path}")
        logger.info(f"IP Cliente: {request.remote_addr}")
        if request.is_json:
            logger.info("Tiene datos JSON: Sí")
        else:
            logger.info("Tiene datos JSON: No")
    except Exception as e:
        # Silenciar errores de logging
        pass

@app.after_request
def after_request(response):
    try:
        logger.info(f"Respuesta enviada: {response.status_code}")
    except Exception as e:
        # Silenciar errores de logging
        pass
    return response

def set_table_borders(table):
    """Aplica bordes negros a toda la tabla"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Crear elemento de bordes
    tblBorders = OxmlElement('w:tblBorders')
    
    # Definir todos los bordes (top, left, bottom, right, insideH, insideV)
    border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    
    for border_type in border_types:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')  # Tipo de línea
        border.set(qn('w:sz'), '4')        # Grosor (4 = 0.5pt)
        border.set(qn('w:space'), '0')     # Espaciado
        border.set(qn('w:color'), '000000') # Color negro
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # Centrar tabla
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def format_table_headers(table):
    """Formatea los headers en negrita y centrados"""
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Inches(0.12)  # 12pt

def format_table_content(table):
    """Formatea el contenido de las tablas"""
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Inches(0.1)  # 10pt

# Configuración de momentos para cada modalidad
MODALIDADES_CONFIG = {
    'abj': ['planteamiento_juego', 'desarrollo_actividades', 'compartamos_experiencia', 'comunidad_juego'],
    'centros': ['contacto_realidad', 'identificacion_integracion', 'expresion'],
    'centros de interes': ['contacto_realidad', 'identificacion_integracion', 'expresion'],
    'centros de interés': ['contacto_realidad', 'identificacion_integracion', 'expresion'],
    'talleres': ['situacion_inicial', 'organizacion_acciones', 'puesta_marcha', 'valoramos_aprendido'],
    'rincones': ['punto_partida', 'asamblea_inicial', 'exploracion_rincones', 'exploracion_descubrimiento', 'compartimos_aprendido', 'evaluamos_experiencia'],
    'rincones de aprendizaje': ['punto_partida', 'asamblea_inicial', 'exploracion_rincones', 'exploracion_descubrimiento', 'compartimos_aprendido', 'evaluamos_experiencia'],
    'proyecto': ['problematizacion', 'desarrollo_proyecto', 'comunicacion', 'integracion', 'reflexion'],
    'unidad': ['lectura_realidad', 'identificacion_trama', 'planificacion', 'exploracion'],
    'unidad didactica': ['lectura_realidad', 'identificacion_trama', 'planificacion', 'exploracion'],
    'unidad didáctica': ['lectura_realidad', 'identificacion_trama', 'planificacion', 'exploracion']
}

@app.route('/', methods=['GET'])
def home():
    """Endpoint raíz para verificar que el servidor está funcionando"""
    logger.info("HOME: Endpoint raíz accedido")
    return jsonify({'status': 'OK', 'message': 'Servidor Plantcher Word está funcionando', 'version': '2.0 - Backend Separado'})

@app.route('/test', methods=['GET'])
def test_connection():
    """Endpoint de prueba para verificar conectividad"""
    logger.info("TEST: Endpoint de prueba accedido")
    return jsonify({'status': 'OK', 'message': 'Servidor funcionando correctamente', 'rutas_disponibles': ['/generar-word', '/test', '/modalidades']})

@app.route('/test-post', methods=['POST'])
def test_post():
    """Endpoint de prueba para verificar solicitudes POST"""
    logger.info("TEST-POST: Endpoint de prueba POST accedido")
    data = request.json
    logger.info(f"Datos recibidos en POST: {data}")
    return jsonify({'status': 'OK', 'message': 'POST funcionando correctamente', 'datos_recibidos': data})

@app.route('/modalidades', methods=['GET'])
def get_modalidades():
    """Endpoint para obtener las modalidades disponibles y sus momentos"""
    logger.info("MODALIDADES: Solicitando lista de modalidades")
    return jsonify({
        'modalidades': list(MODALIDADES_CONFIG.keys()),
        'config': MODALIDADES_CONFIG
    })

@app.route('/generar-word', methods=['POST'])
def generar_word():
    """Generar documento Word con la plantilla ABJ manteniendo estructura original"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No se recibieron datos"}), 400
        
        modalidad = data.get('modalidad', '').lower()
        
        logger.info(f"Modalidad recibida: '{modalidad}'")
        logger.info(f"Modalidades disponibles: {list(MODALIDADES_CONFIG.keys())}")
        
        if modalidad not in MODALIDADES_CONFIG:
            logger.error(f"Modalidad '{modalidad}' no encontrada")
            return jsonify({
                "error": f"Modalidad '{modalidad}' no válida",
                "modalidades_disponibles": list(MODALIDADES_CONFIG.keys())
            }), 400
        
        # Crear documento
        doc = Document()
        
        # Título principal
        title = doc.add_heading(f"Planeación {modalidad.upper()}: {data.get('titulo', '')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")

        # === TABLA 1: DATOS GENERALES ===
        table1 = doc.add_table(rows=2, cols=3)
        
        # Headers
        headers1 = ['Periodo de Aplicación', 'Propósito', 'Relevancia Social']
        for i, header in enumerate(headers1):
            table1.rows[0].cells[i].text = header
        
        # Datos
        row1_data = [
            data.get('periodoAplicacion', ''),
            data.get('proposito', ''),
            data.get('relevanciaSocial', '')
        ]
        for i, dato in enumerate(row1_data):
            table1.rows[1].cells[i].text = dato
        
        set_table_borders(table1)
        format_table_headers(table1)
        format_table_content(table1)
        doc.add_paragraph("")

        # === TABLA 2: CONTENIDO CURRICULAR ===
        campos = data.get('camposFormativos', [])
        contenidos = data.get('contenidos', [])
        procesos = data.get('procesosDesarrollo', [])
        relaciones = data.get('relacionContenidos', {})
        
        max_rows = max(len(campos), len(contenidos), len(procesos))
        table2 = doc.add_table(rows=max_rows + 1, cols=5)
        
        # Headers
        headers2 = ['Campos Formativos', 'Contenidos', 'Procesos de Desarrollo', 'Relación de Contenidos', 'Eje Articulador']
        for i, header in enumerate(headers2):
            table2.rows[0].cells[i].text = header
        
        # Datos
        for i in range(max_rows):
            campo = campos[i] if i < len(campos) else ''
            contenido = contenidos[i] if i < len(contenidos) else ''
            
            # Procesos
            procesos_str = ''
            if procesos and i < len(procesos):
                campo_proc = procesos[i]
                gradosPorContenido = campo_proc.get('gradosPorContenido', {})
                for cont, grados in gradosPorContenido.items():
                    procesos_str += f"{cont}\n"
                    for grado, elementos in grados.items():
                        procesos_str += f"  Grado {grado}:\n"
                        for el in elementos:
                            procesos_str += f"    • {el}\n"
            
            relacion = relaciones.get(campo, '') if campo in relaciones else ''
            eje = data.get('ejeArticulador', '') if i == 0 else ''
            
            table2.rows[i + 1].cells[0].text = campo
            table2.rows[i + 1].cells[1].text = contenido
            table2.rows[i + 1].cells[2].text = procesos_str.strip()
            table2.rows[i + 1].cells[3].text = relacion
            table2.rows[i + 1].cells[4].text = eje
        
        set_table_borders(table2)
        format_table_headers(table2)
        format_table_content(table2)
        doc.add_paragraph("")

        # === TABLA 3: MOMENTOS (Específicos por modalidad) ===
        momentos = data.get('momentos', {})
        momentos_modalidad = MODALIDADES_CONFIG.get(modalidad, ['Inicio', 'Desarrollo', 'Cierre'])
        
        logger.info(f"Momentos recibidos: {list(momentos.keys())}")
        logger.info(f"Momentos esperados para {modalidad}: {momentos_modalidad}")
        
        table3 = doc.add_table(rows=len(momentos_modalidad) + 1, cols=2)
        
        # Headers
        table3.rows[0].cells[0].text = 'Momentos'
        table3.rows[0].cells[1].text = 'Descripción'
        
        # Datos específicos por modalidad
        for idx, momento in enumerate(momentos_modalidad, start=1):
            table3.rows[idx].cells[0].text = momento
            # Buscar la descripción del momento con diferentes variaciones
            descripcion = ''
            # Buscar por nombre exacto
            if momento in momentos:
                descripcion = momentos[momento]
            # Buscar por nombre sin espacios ni mayúsculas
            elif momento.lower().replace(' ', '') in [k.lower().replace(' ', '') for k in momentos.keys()]:
                for k, v in momentos.items():
                    if k.lower().replace(' ', '') == momento.lower().replace(' ', ''):
                        descripcion = v
                        break
            # Si no encuentra, buscar por palabras clave
            elif not descripcion:
                momento_lower = momento.lower()
                for k, v in momentos.items():
                    k_lower = k.lower()
                    if (momento_lower in k_lower) or (k_lower in momento_lower):
                        descripcion = v
                        break
            
            table3.rows[idx].cells[1].text = descripcion
            
            # Log para debugging
            logger.info(f"Momento: '{momento}' -> Descripción: '{descripcion[:50]}...' (encontrada: {bool(descripcion)})")
        
        set_table_borders(table3)
        format_table_headers(table3)
        format_table_content(table3)
        doc.add_paragraph("")

        # === TABLA 4: VARIANTES ===
        table4 = doc.add_table(rows=2, cols=1)
        table4.rows[0].cells[0].text = 'Posibles Variantes'
        table4.rows[1].cells[0].text = data.get('posiblesVariantes', '')
        
        set_table_borders(table4)
        format_table_headers(table4)
        format_table_content(table4)
        doc.add_paragraph("")

        # === TABLA 5: RECURSOS ===
        materiales = data.get('materiales', [])
        espacios = data.get('espacios', [])
        produccion = data.get('produccionSugerida', [])
        
        table5 = doc.add_table(rows=2, cols=3)
        table5.rows[0].cells[0].text = 'Materiales'
        table5.rows[0].cells[1].text = 'Espacios'
        table5.rows[0].cells[2].text = 'Producción Sugerida'
        
        table5.rows[1].cells[0].text = '\n'.join(f'• {m}' for m in materiales)
        table5.rows[1].cells[1].text = '\n'.join(f'• {e}' for e in espacios)
        table5.rows[1].cells[2].text = '\n'.join(f'• {p}' for p in produccion)
        
        set_table_borders(table5)
        format_table_headers(table5)
        format_table_content(table5)

        # Guardar en memoria
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"planeacion_{modalidad}.docx"
        
        logger.info(f"Generando archivo para modalidad: {modalidad}")
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"ERROR: {str(e)}")
        return jsonify({"error": f"Error interno del servidor: {str(e)}"}), 500

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    logger.info("INICIANDO SERVIDOR PLANTCHER WORD - BACKEND SEPARADO")
    logger.info(f"Servidor ejecutándose en puerto: {port}")
    logger.info("Rutas disponibles:")
    logger.info("   POST /generar-word - Generar documento Word")
    logger.info("   GET  /test        - Prueba de conectividad")
    logger.info("   POST /test-post   - Prueba de solicitudes POST")
    logger.info("   GET  /modalidades - Lista de modalidades")
    app.run(debug=False, host='0.0.0.0', port=port)
