"""
Versión del servidor optimizada para Windows/waitress
"""
from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import sys
import logging

app = Flask(__name__)

# Configurar logging para evitar problemas con waitress
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@app.before_request
def log_request_info():
    try:
        logger.info(f"SOLICITUD ENTRANTE: {request.method} {request.path}")
        logger.info(f"IP Cliente: {request.remote_addr}")
        if request.is_json:
            logger.info("Tiene datos JSON: Sí")
        else:
            logger.info("Tiene datos JSON: No")
    except Exception as e:
        # Silenciar errores de logging
        pass

@app.after_request
def after_request(response):
    try:
        logger.info(f"Respuesta enviada: {response.status_code}")
    except Exception as e:
        # Silenciar errores de logging
        pass
    return response

def set_table_borders(table):
    """Establecer bordes para todas las celdas de la tabla"""
    tbl = table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell.append(tcPr)
            
            # Añadir bordes
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

def format_table_headers(table):
    """Formatea los headers en negrita y centrados"""
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Inches(0.12)  # 12pt

def format_table_content(table):
    """Formatea el contenido de las tablas"""
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Inches(0.1)  # 10pt

# Configuración de momentos para cada modalidad
MODALIDADES_CONFIG = {
    'abj': ['Inicio', 'Desarrollo', 'Cierre'],
    'centros': ['Inicio', 'Desarrollo', 'Cierre'],
    'talleres': ['Inicio', 'Desarrollo', 'Cierre'],
    'rincones': ['Inicio', 'Desarrollo', 'Cierre'],
    'proyecto': ['Problematización', 'Desarrollo del Proyecto', 'Comunicación', 'Integración', 'Reflexión'],
    'unidad': ['Inicio', 'Desarrollo', 'Cierre', 'Transversalidad', 'Reflexión', 'Conclusión y Valoración']
}

def crear_tabla_principal():
    """Crear la tabla principal del documento"""
    tabla_principal = [
        ["Campos Formativos", "Contenidos", "Procesos de Desarrollo de Aprendizaje"],
        ["", "", ""],
        ["", "", ""],
        ["", "", ""]
    ]
    return tabla_principal

def crear_tabla_periodo_proposito():
    """Crear tabla de período y propósito"""
    tabla_periodo = [
        ["Período de realización", "Propósito"],
        ["", ""]
    ]
    return tabla_periodo

def crear_tabla_momentos(modalidad):
    """Crear tabla de momentos según la modalidad"""
    momentos = MODALIDADES_CONFIG.get(modalidad, ['Inicio', 'Desarrollo', 'Cierre'])
    
    tabla_momentos = [["Momentos", "Situaciones de Aprendizaje", "Recursos", "Tiempo"]]
    
    for momento in momentos:
        tabla_momentos.append([momento, "", "", ""])
    
    return tabla_momentos

def crear_tabla_evaluacion():
    """Crear tabla de evaluación"""
    tabla_evaluacion = [
        ["Aspectos a Evaluar", "Técnicas", "Instrumentos"],
        ["", "", ""],
        ["", "", ""],
        ["", "", ""]
    ]
    return tabla_evaluacion

def crear_tabla_bibliografia():
    """Crear tabla de bibliografía"""
    tabla_bibliografia = [
        ["Bibliografía"],
        [""],
        [""],
        [""]
    ]
    return tabla_bibliografia

@app.route('/test', methods=['GET'])
def test():
    """Endpoint de prueba para verificar conectividad"""
    return jsonify({
        "status": "success",
        "message": "Servidor Plantcher Word funcionando correctamente",
        "version": "2.0 - Optimizado para Windows"
    })

@app.route('/test-post', methods=['POST'])
def test_post():
    """Endpoint de prueba para solicitudes POST"""
    data = request.get_json() if request.is_json else {}
    return jsonify({
        "status": "success",
        "message": "Solicitud POST recibida correctamente",
        "received_data": data
    })

@app.route('/modalidades', methods=['GET'])
def get_modalidades():
    """Obtener lista de modalidades disponibles"""
    return jsonify({
        "modalidades": list(MODALIDADES_CONFIG.keys()),
        "descripcion": {
            "abj": "Aprendizaje Basado en Juegos",
            "centros": "Centros de Aprendizaje",
            "talleres": "Talleres",
            "rincones": "Rincones de Aprendizaje",
            "proyecto": "Proyecto",
            "unidad": "Unidad Didáctica"
        }
    })

@app.route('/generar-word', methods=['POST'])
def generar_word():
    """Generar documento Word con la plantilla ABJ"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No se recibieron datos"}), 400
        
        modalidad = data.get('modalidad', '').lower()
        
        if modalidad not in MODALIDADES_CONFIG:
            return jsonify({
                "error": f"Modalidad '{modalidad}' no válida",
                "modalidades_disponibles": list(MODALIDADES_CONFIG.keys())
            }), 400
        
        # Crear documento
        doc = Document()
        
        # Título principal
        titulo = doc.add_heading('PLANTILLA DIDÁCTICA', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo con modalidad
        modalidad_upper = modalidad.upper()
        subtitulo = doc.add_heading(f'{modalidad_upper}', 1)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Tabla Principal (Campos Formativos, Contenidos, Procesos)
        doc.add_heading('1. INFORMACIÓN CURRICULAR', 2)
        tabla_principal_data = crear_tabla_principal()
        tabla_principal = doc.add_table(rows=len(tabla_principal_data), cols=len(tabla_principal_data[0]))
        tabla_principal.style = 'Table Grid'
        tabla_principal.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(tabla_principal_data):
            for j, cell_data in enumerate(row_data):
                tabla_principal.cell(i, j).text = cell_data
        
        set_table_borders(tabla_principal)
        format_table_headers(tabla_principal)
        format_table_content(tabla_principal)
        
        doc.add_paragraph()
        
        # 2. Tabla de Período y Propósito
        doc.add_heading('2. PERÍODO Y PROPÓSITO', 2)
        tabla_periodo_data = crear_tabla_periodo_proposito()
        tabla_periodo = doc.add_table(rows=len(tabla_periodo_data), cols=len(tabla_periodo_data[0]))
        tabla_periodo.style = 'Table Grid'
        tabla_periodo.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(tabla_periodo_data):
            for j, cell_data in enumerate(row_data):
                tabla_periodo.cell(i, j).text = cell_data
        
        set_table_borders(tabla_periodo)
        format_table_headers(tabla_periodo)
        format_table_content(tabla_periodo)
        
        doc.add_paragraph()
        
        # 3. Tabla de Momentos (específica para cada modalidad)
        doc.add_heading('3. MOMENTOS DIDÁCTICOS', 2)
        tabla_momentos_data = crear_tabla_momentos(modalidad)
        tabla_momentos = doc.add_table(rows=len(tabla_momentos_data), cols=len(tabla_momentos_data[0]))
        tabla_momentos.style = 'Table Grid'
        tabla_momentos.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(tabla_momentos_data):
            for j, cell_data in enumerate(row_data):
                tabla_momentos.cell(i, j).text = cell_data
        
        set_table_borders(tabla_momentos)
        format_table_headers(tabla_momentos)
        format_table_content(tabla_momentos)
        
        doc.add_paragraph()
        
        # 4. Tabla de Evaluación
        doc.add_heading('4. EVALUACIÓN', 2)
        tabla_evaluacion_data = crear_tabla_evaluacion()
        tabla_evaluacion = doc.add_table(rows=len(tabla_evaluacion_data), cols=len(tabla_evaluacion_data[0]))
        tabla_evaluacion.style = 'Table Grid'
        tabla_evaluacion.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(tabla_evaluacion_data):
            for j, cell_data in enumerate(row_data):
                tabla_evaluacion.cell(i, j).text = cell_data
        
        set_table_borders(tabla_evaluacion)
        format_table_headers(tabla_evaluacion)
        format_table_content(tabla_evaluacion)
        
        doc.add_paragraph()
        
        # 5. Tabla de Bibliografía
        doc.add_heading('5. BIBLIOGRAFÍA', 2)
        tabla_bibliografia_data = crear_tabla_bibliografia()
        tabla_bibliografia = doc.add_table(rows=len(tabla_bibliografia_data), cols=len(tabla_bibliografia_data[0]))
        tabla_bibliografia.style = 'Table Grid'
        tabla_bibliografia.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(tabla_bibliografia_data):
            for j, cell_data in enumerate(row_data):
                tabla_bibliografia.cell(i, j).text = cell_data
        
        set_table_borders(tabla_bibliografia)
        format_table_headers(tabla_bibliografia)
        format_table_content(tabla_bibliografia)
        
        # Guardar en memoria
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"plantilla_{modalidad}.docx"
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error generando documento: {str(e)}")
        return jsonify({"error": f"Error interno del servidor: {str(e)}"}), 500

if __name__ == '__main__':
    from waitress import serve
    import os
    port = int(os.environ.get('PORT', 5000))
    logger.info("🚀 SERVIDOR PLANTCHER WORD - VERSIÓN WINDOWS")
    logger.info(f"📍 Puerto: {port}")
    serve(app, host='0.0.0.0', port=port)
