from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import io
import os

app = Flask(__name__)

def add_table_borders(table):
    """Añade bordes a una tabla de Word"""
    tbl = table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

def create_basic_info_table(doc, data):
    """Crea la tabla de información básica común a todas las modalidades"""
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Datos básicos
    info_rows = [
        ['Modalidad:', data.get('modalidad', '')],
        ['Experiencia:', data.get('experiencia', '')],
        ['Nombre de la Experiencia:', data.get('nombre_experiencia', '')],
        ['Edad:', data.get('edad', '')],
        ['Duración:', data.get('duracion', '')],
        ['Número de Niños/as:', data.get('numero_ninos', '')],
        ['Agente Educativo:', data.get('agente_educativo', '')]
    ]
    
    for i, (label, value) in enumerate(info_rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
        
        # Formatear celdas
        for j in range(2):
            cell = table.cell(i, j)
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            if j == 0:  # Etiquetas en negrita
                cell.paragraphs[0].runs[0].font.bold = True
    
    add_table_borders(table)
    return table

def create_momentos_table(doc, momentos_data, momentos_labels):
    """Crea la tabla de momentos específica para cada modalidad"""
    table = doc.add_table(rows=len(momentos_labels) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezado
    header_row = table.rows[0]
    header_row.cells[0].text = "Momento"
    header_row.cells[1].text = "Descripción"
    
    for cell in header_row.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Momentos
    for i, (label, key) in enumerate(momentos_labels, 1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(momentos_data.get(key, ''))
        
        # Formatear celdas
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(11)
        table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(11)
    
    add_table_borders(table)
    return table

def create_materials_table(doc, data):
    """Crea la tabla de materiales, espacios y producción"""
    doc.add_paragraph()
    
    # Título
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Materiales, Espacios y Producción Sugerida")
    title_run.font.bold = True
    title_run.font.size = Pt(14)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tabla
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Datos de materiales
    materials_rows = [
        ['Materiales:', data.get('materiales', '')],
        ['Espacios:', data.get('espacios', '')],
        ['Producción Sugerida:', data.get('produccion_sugerida', '')],
        ['Variantes:', data.get('variantes', '')]
    ]
    
    for i, (label, value) in enumerate(materials_rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
        
        # Formatear celdas
        table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(11)
        table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(11)
    
    add_table_borders(table)
    return table

# Configuración de momentos para cada modalidad
MODALIDADES_CONFIG = {
    'ABJ': [
        ('1. Planteamiento del Juego', 'planteamiento_juego'),
        ('2. Desarrollo de las Actividades', 'desarrollo_actividades'),
        ('3. Compartamos la Experiencia', 'compartamos_experiencia'),
        ('4. Comunidad de Juego', 'comunidad_juego')
    ],
    'Aprendizaje Basado en Juegos': [
        ('1. Planteamiento del Juego', 'planteamiento_juego'),
        ('2. Desarrollo de las Actividades', 'desarrollo_actividades'),
        ('3. Compartamos la Experiencia', 'compartamos_experiencia'),
        ('4. Comunidad de Juego', 'comunidad_juego')
    ],
    'Centros de Interés': [
        ('1. En contacto de la realidad', 'contacto_realidad'),
        ('2. Identificación e integración', 'identificacion_integracion'),
        ('3. Expresión', 'expresion')
    ],
    'Proyecto': [
        ('1. Punto de partida', 'punto_partida'),
        ('2. Planeación', 'planeacion'),
        ('3. ¡A trabajar!', 'a_trabajar'),
        ('4. Comunicamos nuestros logros', 'comunicamos_logros'),
        ('5. Reflexión sobre el aprendizaje', 'reflexion_aprendizaje')
    ],
    'Rincones de Aprendizaje': [
        ('1. Punto de partida (Saberes previos)', 'punto_partida'),
        ('2. Asamblea inicial y planeación', 'asamblea_inicial'),
        ('3. Exploración de los rincones', 'exploracion_rincones'),
        ('4. Exploración y descubrimiento', 'exploracion_descubrimiento'),
        ('5. Compartimos lo aprendido', 'compartimos_aprendido'),
        ('6. Evaluamos la experiencia', 'evaluamos_experiencia')
    ],
    'Taller Crítico': [
        ('1. Situación inicial', 'situacion_inicial'),
        ('2. Organización de las acciones', 'organizacion_acciones'),
        ('3. Puesta en marcha', 'puesta_marcha'),
        ('4. Valoramos lo aprendido', 'valoramos_aprendido')
    ],
    'Unidad Didáctica': [
        ('1. Lectura de la realidad', 'lectura_realidad'),
        ('2. Identificación de la trama y complejidad', 'identificacion_trama'),
        ('3. Planificación y organización del trabajo', 'planificacion'),
        ('4. Exploración y descubrimiento', 'exploracion'),
        ('5. Participación activa y horizontal', 'participacion'),
        ('6. Conclusión de la experiencia (Valoración)', 'conclusion')
    ]
}

@app.route('/generar-word', methods=['POST'])
def generar_word():
    try:
        # Recibir datos
        data = request.json
        modalidad = data.get('modalidad', '')
        
        # Verificar que la modalidad esté soportada
        if modalidad not in MODALIDADES_CONFIG:
            return jsonify({'error': f'Modalidad "{modalidad}" no soportada'}), 400
        
        # Crear documento
        doc = Document()
        
        # Título principal
        title = doc.add_heading(f'Planeación Pedagógica - {modalidad}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo con nombre de experiencia
        if data.get('nombre_experiencia'):
            subtitle = doc.add_heading(data.get('nombre_experiencia'), level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Tabla de información básica
        create_basic_info_table(doc, data)
        
        doc.add_paragraph()
        
        # Título de momentos
        moments_title = doc.add_paragraph()
        moments_run = moments_title.add_run("Momentos de la Experiencia")
        moments_run.font.bold = True
        moments_run.font.size = Pt(14)
        moments_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tabla de momentos específicos
        momentos_data = data.get('momentos', {})
        momentos_config = MODALIDADES_CONFIG[modalidad]
        create_momentos_table(doc, momentos_data, momentos_config)
        
        # Tabla de materiales
        create_materials_table(doc, data)
        
        # Guardar en memoria
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Nombre del archivo
        nombre_archivo = f"planificacion_{modalidad.lower().replace(' ', '_')}_{data.get('nombre_experiencia', 'experiencia').replace(' ', '_')}.docx"
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=nombre_archivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/modalidades', methods=['GET'])
def get_modalidades():
    """Endpoint para obtener las modalidades disponibles y sus momentos"""
    return jsonify({
        'modalidades': list(MODALIDADES_CONFIG.keys()),
        'config': MODALIDADES_CONFIG
    })

@app.route('/health', methods=['GET'])
def health_check():
    """Endpoint de verificación de salud del servidor"""
    return jsonify({'status': 'OK', 'message': 'Servidor funcionando correctamente'})

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
