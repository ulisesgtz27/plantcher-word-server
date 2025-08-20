from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import logging

app = Flask(__name__)

# Configurar logging para evitar problemas con WSGI servers
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@app.before_request
def log_request_info():
    try:
        logger.info(f"SOLICITUD ENTRANTE: {request.method} {request.path}")
        logger.info(f"IP Cliente: {request.remote_addr}")
        if request.is_json:
            logger.info("Tiene datos JSON: Sí")
        else:
            logger.info("Tiene datos JSON: No")
    except Exception as e:
        # Silenciar errores de logging
        pass

@app.after_request
def after_request(response):
    try:
        logger.info(f"Respuesta enviada: {response.status_code}")
    except Exception as e:
        # Silenciar errores de logging
        pass
    return response

def set_table_borders(table):
    """Aplica bordes negros a toda la tabla"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Crear elemento de bordes
    tblBorders = OxmlElement('w:tblBorders')
    
    # Definir todos los bordes (top, left, bottom, right, insideH, insideV)
    border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    
    for border_type in border_types:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')  # Tipo de línea
        border.set(qn('w:sz'), '4')        # Grosor (4 = 0.5pt)
        border.set(qn('w:space'), '0')     # Espaciado
        border.set(qn('w:color'), '000000') # Color negro
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # Centrar tabla
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def format_table_headers(table):
    """Formatea los headers en negrita y centrados"""
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Inches(0.12)  # 12pt

def format_table_content(table):
    """Formatea el contenido de las tablas"""
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Inches(0.1)  # 10pt

# Configuración de momentos para cada modalidad
MODALIDADES_CONFIG = {
    'ABJ': [
        ('Planteamiento del Juego', 'planteamiento_juego'),
        ('Desarrollo de las Actividades', 'desarrollo_actividades'),
        ('Compartamos la Experiencia', 'compartamos_experiencia'),
        ('Comunidad de Juego', 'comunidad_juego')
    ],
    'Aprendizaje Basado en Juegos': [
        ('Planteamiento del Juego', 'planteamiento_juego'),
        ('Desarrollo de las Actividades', 'desarrollo_actividades'),
        ('Compartamos la Experiencia', 'compartamos_experiencia'),
        ('Comunidad de Juego', 'comunidad_juego')
    ],
    'Centros de Interés': [
        ('En contacto de la realidad', 'contacto_realidad'),
        ('Identificación e integración', 'identificacion_integracion'),
        ('Expresión', 'expresion')
    ],
    'Proyecto': [
        ('Punto de partida', 'punto_partida'),
        ('Planeación', 'planeacion'),
        ('¡A trabajar!', 'a_trabajar'),
        ('Comunicamos nuestros logros', 'comunicamos_logros'),
        ('Reflexión sobre el aprendizaje', 'reflexion_aprendizaje')
    ],
    'Rincones de Aprendizaje': [
        ('Punto de partida (Saberes previos)', 'punto_partida'),
        ('Asamblea inicial y planeación', 'asamblea_inicial'),
        ('Exploración de los rincones', 'exploracion_rincones'),
        ('Exploración y descubrimiento', 'exploracion_descubrimiento'),
        ('Compartimos lo aprendido', 'compartimos_aprendido'),
        ('Evaluamos la experiencia', 'evaluamos_experiencia')
    ],
    'Taller Crítico': [
        ('Situación inicial', 'situacion_inicial'),
        ('Organización de las acciones', 'organizacion_acciones'),
        ('Puesta en marcha', 'puesta_marcha'),
        ('Valoramos lo aprendido', 'valoramos_aprendido')
    ],
    'Unidad Didáctica': [
        ('Lectura de la realidad', 'lectura_realidad'),
        ('Identificación de la trama y complejidad', 'identificacion_trama'),
        ('Planificación y organización del trabajo', 'planificacion'),
        ('Exploración y descubrimiento', 'exploracion'),
        ('Participación activa y horizontal', 'participacion'),
        ('Conclusión de la experiencia (Valoración)', 'conclusion')
    ]
}

@app.route('/generar-word', methods=['POST'])
def generar_word():
    print("🔥 SOLICITUD RECIBIDA EN /generar-word")
    print(f"🌐 Método: {request.method}")
    print(f"📍 Ruta: {request.path}")
    print(f"🔗 URL completa: {request.url}")
    print(f"📋 Headers: {dict(request.headers)}")
    
    try:
        data = request.json
        print(f"📦 Datos recibidos: {data}")
        
        modalidad = data.get('modalidad', '')
        print(f"🎯 Modalidad: {modalidad}")
        
        # Buscar modalidad de forma case-insensitive
        modalidad_encontrada = None
        for config_modalidad in MODALIDADES_CONFIG.keys():
            if modalidad.lower() == config_modalidad.lower():
                modalidad_encontrada = config_modalidad
                break
        
        # Verificar que la modalidad esté soportada
        if modalidad_encontrada is None:
            print(f"❌ Modalidad no soportada: {modalidad}")
            print(f"📋 Modalidades disponibles: {list(MODALIDADES_CONFIG.keys())}")
            return jsonify({'error': f'Modalidad "{modalidad}" no soportada'}), 400
        
        print(f"✅ Modalidad encontrada: {modalidad_encontrada}")
        
        # Usar la modalidad encontrada para el resto del proceso
        modalidad = modalidad_encontrada

        doc = Document()
        
        # Título principal
        title = doc.add_heading(f"Planeación {modalidad}: {data.get('titulo', '')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")

        # === TABLA 1: DATOS GENERALES ===
        table1 = doc.add_table(rows=2, cols=3)
        
        # Headers
        headers1 = ['Periodo de Aplicación', 'Propósito', 'Relevancia Social']
        for i, header in enumerate(headers1):
            table1.rows[0].cells[i].text = header
        
        # Datos
        row1_data = [
            data.get('periodoAplicacion', ''),
            data.get('proposito', ''),
            data.get('relevanciaSocial', '')
        ]
        for i, dato in enumerate(row1_data):
            table1.rows[1].cells[i].text = dato
        
        set_table_borders(table1)
        format_table_headers(table1)
        format_table_content(table1)
        doc.add_paragraph("")

        # === TABLA 2: CONTENIDO CURRICULAR ===
        campos = data.get('camposFormativos', [])
        contenidos = data.get('contenidos', [])
        procesos = data.get('procesosDesarrollo', [])
        relaciones = data.get('relacionContenidos', {})
        
        max_rows = max(len(campos), len(contenidos), len(procesos))
        table2 = doc.add_table(rows=max_rows + 1, cols=5)
        
        # Headers
        headers2 = ['Campos Formativos', 'Contenidos', 'Procesos de Desarrollo', 'Relación de Contenidos', 'Eje Articulador']
        for i, header in enumerate(headers2):
            table2.rows[0].cells[i].text = header
        
        # Datos
        for i in range(max_rows):
            campo = campos[i] if i < len(campos) else ''
            contenido = contenidos[i] if i < len(contenidos) else ''
            
            # Procesos
            procesos_str = ''
            if procesos and i < len(procesos):
                campo_proc = procesos[i]
                gradosPorContenido = campo_proc.get('gradosPorContenido', {})
                for cont, grados in gradosPorContenido.items():
                    procesos_str += f"{cont}\n"
                    for grado, elementos in grados.items():
                        procesos_str += f"  Grado {grado}:\n"
                        for el in elementos:
                            procesos_str += f"    • {el}\n"
            
            relacion = relaciones.get(campo, '') if campo in relaciones else ''
            eje = data.get('ejeArticulador', '') if i == 0 else ''
            
            table2.rows[i + 1].cells[0].text = campo
            table2.rows[i + 1].cells[1].text = contenido
            table2.rows[i + 1].cells[2].text = procesos_str.strip()
            table2.rows[i + 1].cells[3].text = relacion
            table2.rows[i + 1].cells[4].text = eje
        
        set_table_borders(table2)
        format_table_headers(table2)
        format_table_content(table2)
        doc.add_paragraph("")

        # === TABLA 3: MOMENTOS (ADAPTADA SEGÚN MODALIDAD) ===
        momentos = data.get('momentos', {})
        momentos_config = MODALIDADES_CONFIG[modalidad]
        table3 = doc.add_table(rows=len(momentos_config) + 1, cols=2)
        
        # Headers
        table3.rows[0].cells[0].text = 'Momentos'
        table3.rows[0].cells[1].text = 'Descripción'
        
        # Momentos específicos según modalidad
        for idx, (label, key) in enumerate(momentos_config, start=1):
            table3.rows[idx].cells[0].text = label
            table3.rows[idx].cells[1].text = momentos.get(key, '')
        
        set_table_borders(table3)
        format_table_headers(table3)
        format_table_content(table3)
        doc.add_paragraph("")

        # === TABLA 4: VARIANTES ===
        table4 = doc.add_table(rows=2, cols=1)
        table4.rows[0].cells[0].text = 'Posibles Variantes'
        table4.rows[1].cells[0].text = data.get('posiblesVariantes', '')
        
        set_table_borders(table4)
        format_table_headers(table4)
        format_table_content(table4)
        doc.add_paragraph("")

        # === TABLA 5: RECURSOS ===
        materiales = data.get('materiales', [])
        espacios = data.get('espacios', [])
        produccion = data.get('produccionSugerida', [])
        
        table5 = doc.add_table(rows=2, cols=3)
        table5.rows[0].cells[0].text = 'Materiales'
        table5.rows[0].cells[1].text = 'Espacios'
        table5.rows[0].cells[2].text = 'Producción Sugerida'
        
        table5.rows[1].cells[0].text = '\n'.join(f'• {m}' for m in materiales)
        table5.rows[1].cells[1].text = '\n'.join(f'• {e}' for e in espacios)
        table5.rows[1].cells[2].text = '\n'.join(f'• {p}' for p in produccion)
        
        set_table_borders(table5)
        format_table_headers(table5)
        format_table_content(table5)

        # Guardar en memoria y enviar
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        logger.info(f"Generando archivo para modalidad: {modalidad}")
        
        return send_file(
            f,
            as_attachment=True,
            download_name=f"planeacion_{modalidad.lower().replace(' ', '_')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"ERROR: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/', methods=['GET'])
def home():
    """Endpoint raíz para verificar que el servidor está funcionando"""
    logger.info("HOME: Endpoint raíz accedido")
    return jsonify({'status': 'OK', 'message': 'Servidor Plantcher Word está funcionando', 'version': '1.0'})

@app.route('/test', methods=['GET'])
def test_connection():
    """Endpoint de prueba para verificar conectividad"""
    logger.info("TEST: Endpoint de prueba accedido")
    return jsonify({'status': 'OK', 'message': 'Servidor funcionando correctamente', 'rutas_disponibles': ['/generar-word', '/test', '/modalidades']})

@app.route('/test-post', methods=['POST'])
def test_post():
    """Endpoint de prueba para verificar solicitudes POST"""
    logger.info("TEST-POST: Endpoint de prueba POST accedido")
    data = request.json
    logger.info(f"Datos recibidos en POST: {data}")
    return jsonify({'status': 'OK', 'message': 'POST funcionando correctamente', 'datos_recibidos': data})

@app.route('/modalidades', methods=['GET'])
def get_modalidades():
    """Endpoint para obtener las modalidades disponibles y sus momentos"""
    logger.info("MODALIDADES: Solicitando lista de modalidades")
    return jsonify({
        'modalidades': list(MODALIDADES_CONFIG.keys()),
        'config': MODALIDADES_CONFIG
    })

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    logger.info("INICIANDO SERVIDOR PLANTCHER WORD")
    logger.info(f"Servidor ejecutándose en puerto: {port}")
    logger.info("Rutas disponibles:")
    logger.info("   POST /generar-word - Generar documento Word")
    logger.info("   GET  /test        - Prueba de conectividad")
    logger.info("   POST /test-post   - Prueba de solicitudes POST")
    logger.info("   GET  /modalidades - Lista de modalidades")
    app.run(debug=False, host='0.0.0.0', port=port)
